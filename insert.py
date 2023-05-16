import os
import re

import docx
from db import extract_db
from docx.shared import Inches


def add_width(
    doc: docx.document.Document, fname: str, width: float, position: int
) -> None:
    if os.path.exists(fname):
        doc.paragraphs[position + 2].add_run().add_picture(fname, width=Inches(width))
        print(fname, "is added")
    else:
        print(fname, "cannot be found!")


def add_height(
    doc: docx.document.Document, fname: str, height: float, position: int
) -> None:
    if os.path.exists(fname):
        doc.paragraphs[position + 2].add_run().add_picture(fname, height=Inches(height))
        print(fname, "is added")
    else:
        print(fname, "cannot be found!")


def filter_by_document(docFile_path: str) -> tuple:
    paragraph_index = []  # index that contain the words/phrase
    file_name = []
    extension = []
    prompt = []
    file_number = []  # for paragraph_index
    priority = []
    priority_value = []
    data = extract_db()

    doc = docx.Document(docFile_path)

    for i, para in enumerate(doc.paragraphs):
        for k in range(len(data)):
            if data[k][2] in para.text:
                paragraph_index.append(i)
                file_name.append(data[k][0])
                extension.append(data[k][1])
                prompt.append(data[k][2])
                file_number.append(data[k][3])
                priority.append(data[k][4])
                priority_value.append(data[k][5])

    return (
        paragraph_index,
        file_name,
        extension,
        prompt,
        file_number,
        priority,
        priority_value,
    )


def duplicate_check(x: list) -> bool:
    for element in x:
        if x.count(element) > 1:
            return True
    return False


def auto_insert(
    docFile_path: str,
    imgFolder_path: str,
    document_instruction: tuple,
    image_instruction: list,
    file_name: list,
) -> None:
    doc = docx.Document(docFile_path)

    image_name = []
    for i in range(len(file_name)):
        if not (i in image_instruction):
            image_name.append(file_name[i])
    print(image_name)

    multiple_file = []
    multiple_count = []
    for j in range(len(document_instruction[0])):
        if document_instruction[4][j] == "Single":
            if document_instruction[5][j] == "Height":
                fname = "".join(
                    [
                        imgFolder_path,
                        "/",
                        document_instruction[1][j],
                        document_instruction[2][j],
                    ]
                )
                add_height(
                    doc=doc,
                    fname=fname,
                    height=document_instruction[6][j],
                    position=document_instruction[0][j],
                )
            elif document_instruction[5][j] == "Width":
                fname = "".join(
                    [
                        imgFolder_path,
                        "/",
                        document_instruction[1][j],
                        document_instruction[2][j],
                    ]
                )
                add_width(
                    doc=doc,
                    fname=fname,
                    width=document_instruction[6][j],
                    position=document_instruction[0][j],
                )
        elif document_instruction[4][j] == "Multiple":
            if any(document_instruction[1][j] == x for x in multiple_file):
                multiple_count[multiple_file.index(document_instruction[1][j])] += 1
            else:
                multiple_file.append(document_instruction[1][j])
                multiple_count.append(1)
            fname = "".join(
                [
                    imgFolder_path,
                    "/",
                    document_instruction[1][j],
                    "_",
                    str(
                        multiple_count[multiple_file.index(document_instruction[1][j])]
                    ),
                    document_instruction[2][j],
                ]
            )
            if document_instruction[5][j] == "Height":
                add_height(
                    doc=doc,
                    fname=fname,
                    height=document_instruction[6][j],
                    position=document_instruction[0][j],
                )
            elif document_instruction[5][j] == "Width":
                add_width(
                    doc=doc,
                    fname=fname,
                    width=document_instruction[6][j],
                    position=document_instruction[0][j],
                )
    doc.save(docFile_path)
